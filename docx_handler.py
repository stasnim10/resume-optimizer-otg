"""
Handle .docx file operations: extraction and deterministic replacement
Word-style Find & Replace that preserves formatting
"""
from docx import Document
from typing import List, Tuple, Dict, Any


def extract_text(doc_path: str) -> str:
    """
    Extract plain text from .docx file.
    
    Args:
        doc_path: Path to .docx file
        
    Returns:
        Full plain text from document
    """
    doc = Document(doc_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


def replace_paragraph_text(para, replacement: str) -> bool:
    """
    Replace entire paragraph text while preserving formatting.
    
    Args:
        para: Paragraph object
        replacement: New text for the entire paragraph
        
    Returns:
        True if replacement successful
    """
    if not para.runs:
        para.add_run(replacement)
        return True
    
    # Get formatting from first run
    first_run = para.runs[0]
    first_fmt = {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'underline': first_run.underline,
        'font_name': first_run.font.name,
        'font_size': first_run.font.size,
        'font_color': first_run.font.color.rgb if hasattr(first_run.font.color, 'rgb') else None,
    }
    
    # Clear all runs
    for run in para.runs:
        r = run._element
        r.getparent().remove(r)
    
    # Add replacement text with original formatting
    new_run = para.add_run(replacement)
    new_run.bold = first_fmt['bold']
    new_run.italic = first_fmt['italic']
    new_run.underline = first_fmt['underline']
    if first_fmt['font_name']:
        new_run.font.name = first_fmt['font_name']
    if first_fmt['font_size']:
        new_run.font.size = first_fmt['font_size']
    if first_fmt['font_color']:
        try:
            new_run.font.color.rgb = first_fmt['font_color']
        except:
            pass
    
    return True


def replace_exact_paragraph(doc, anchor: str, new_text: str) -> str:
    """
    Replace a paragraph using STRICT FULL-TEXT EQUALITY.
    
    This is deterministic and safe:
    - Matches only if paragraph.text.strip() == anchor.strip()
    - No substring matching
    - No partial anchors
    - Immune to formatting differences
    
    Args:
        doc: Document object
        anchor: FULL paragraph text (exact match required)
        new_text: Replacement text
        
    Returns:
        Success message
        
    Raises:
        ValueError: If no match or duplicates found
    """
    matches = []
    for idx, para in enumerate(doc.paragraphs):
        # STRICT EQUALITY - not substring matching
        if para.text.strip() == anchor.strip():
            matches.append((idx, para))
    
    if len(matches) == 0:
        raise ValueError(f"❌ Anchor not found (must be FULL paragraph text): '{anchor[:80]}...'")
    elif len(matches) > 1:
        raise ValueError(f"❌ Multiple matches found ({len(matches)}) for anchor: '{anchor[:80]}...'")
    
    idx, para = matches[0]
    
    # Replace entire paragraph
    replace_paragraph_text(para, new_text)
    
    return f"✅ Replaced: '{anchor[:60]}...'"


def apply_replacements(doc_path: str, payload: Dict[str, Any]) -> Tuple[bool, str]:
    """
    Apply all replacements from payload to document.
    Preserves formatting like Word's Find & Replace.
    
    Args:
        doc_path: Path to .docx file
        payload: Validated replacement payload
        
    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        doc = Document(doc_path)
        replaced_anchors = []
        errors = []
        
        # Summary replacement
        if "summary_replacement" in payload:
            sr = payload["summary_replacement"]
            try:
                msg = replace_exact_paragraph(doc, sr["match_anchor"], sr["replacement_text"])
                replaced_anchors.append(sr["match_anchor"])
                print(msg)
            except ValueError as e:
                errors.append(str(e))
                print(str(e))
        
        # Bullet replacements
        if "bullet_replacements" in payload:
            for idx, bullet in enumerate(payload["bullet_replacements"]):
                try:
                    msg = replace_exact_paragraph(doc, bullet["match_anchor"], bullet["replacement_text"])
                    replaced_anchors.append(bullet["match_anchor"])
                    print(msg)
                except ValueError as e:
                    errors.append(str(e))
                    print(str(e))
        
        # If any errors, fail
        if errors:
            error_summary = "\n".join(errors)
            return False, f"Errors during replacement:\n{error_summary}"
        
        # Save file
        output_path = generate_output_filename(doc_path)
        doc.save(output_path)
        
        msg = f"✅ Successfully optimized resume!\nReplaced {len(replaced_anchors)} section(s).\n📄 Saved: {output_path}"
        return True, msg
        
    except Exception as e:
        return False, f"❌ Error: {str(e)}"


def generate_output_filename(original_path: str) -> str:
    """
    Generate output filename with _Optimized suffix.
    
    Args:
        original_path: Original file path
        
    Returns:
        New file path with _Optimized suffix
    """
    from pathlib import Path
    path = Path(original_path)
    output_name = f"{path.stem}_Optimized{path.suffix}"
    output_path = path.parent / output_name
    return str(output_path)


